"""
tools/word_exporter.py
----------------------
Exports research reports to Microsoft Word (.docx) format.
Uses python-docx for professional document formatting.
Supports the same markdown structure as the PDF exporter.
"""

import os
import re
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


def _set_heading_color(paragraph, r: int, g: int, b: int):
    """Helper to set heading color via run formatting."""
    for run in paragraph.runs:
        run.font.color.rgb = RGBColor(r, g, b)


def _add_horizontal_line(doc):
    """Add a horizontal rule to the document."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '6366f1')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def generate_word(
    topic: str,
    report: str,
    quality_score: int,
    verified_count: int,
    llm_type: str = "groq"
) -> str:
    """
    Generate a professional Word document from the research report.

    Args:
        topic: Research topic string
        report: Full markdown report text
        quality_score: Critic quality score (0-100)
        verified_count: Number of verified claims
        llm_type: LLM provider used

    Returns:
        Path to the saved .docx file.

    Raises:
        ImportError: If python-docx is not installed.
        RuntimeError: If document generation fails.
    """
    if not DOCX_AVAILABLE:
        raise ImportError(
            "python-docx is not installed. Run: pip install python-docx"
        )

    os.makedirs("output", exist_ok=True)
    clean_topic = "".join(c for c in topic[:40] if c.isalnum() or c in (' ', '_', '-')).strip()
    clean_topic = clean_topic.replace(' ', '_')
    filename = f"output/research_{clean_topic}_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"

    doc = Document()

    # ── Page margins ──────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    # ── Cover / Header ────────────────────────────────────────────────────────
    # System title
    sys_title = doc.add_paragraph()
    sys_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sys_title.add_run("Sage — Multi-Agent Research Intelligence System")
    run.font.size   = Pt(11)
    run.font.color.rgb = RGBColor(0x63, 0x66, 0xF1)  # Indigo
    run.font.bold   = True
    sys_title.paragraph_format.space_after = Pt(4)

    # Main title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(f"Research Report")
    run.font.size  = Pt(26)
    run.font.bold  = True
    run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
    title_para.paragraph_format.space_after = Pt(6)

    # Topic
    topic_para = doc.add_paragraph()
    topic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = topic_para.add_run(topic)
    run.font.size  = Pt(14)
    run.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
    run.font.italic = True
    topic_para.paragraph_format.space_after = Pt(16)

    _add_horizontal_line(doc)

    # Metadata table
    meta_table = doc.add_table(rows=5, cols=2)
    meta_table.style = 'Table Grid'
    meta_data = [
        ("Generated On",       datetime.now().strftime("%B %d, %Y at %H:%M")),
        ("Research Topic",     topic),
        ("Quality Score",      f"{quality_score}/100"),
        ("Verified Claims",    str(verified_count)),
        ("AI System",          f"LangGraph Multi-Agent • {llm_type.upper()}"),
    ]
    for i, (label, value) in enumerate(meta_data):
        row = meta_table.rows[i]
        label_cell = row.cells[0]
        value_cell = row.cells[1]
        label_cell.text = label
        value_cell.text = value
        # Style label cell
        for para in label_cell.paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)

    doc.add_paragraph()  # Spacer

    # ── Report Content Parser ──────────────────────────────────────────────────
    lines = report.split('\n')
    i = 0
    while i < len(lines):
        raw = lines[i]
        line = raw.strip()
        i += 1

        if not line:
            continue

        # H1
        if line.startswith('# '):
            h = doc.add_heading(line[2:], level=1)
            _set_heading_color(h, 0x1E, 0x3A, 0x8A)  # Dark blue
            h.paragraph_format.space_before = Pt(14)

        # H2
        elif line.startswith('## '):
            h = doc.add_heading(line[3:], level=2)
            _set_heading_color(h, 0x0F, 0x76, 0x6E)  # Teal
            h.paragraph_format.space_before = Pt(10)

        # H3
        elif line.startswith('### '):
            h = doc.add_heading(line[4:], level=3)
            _set_heading_color(h, 0x63, 0x66, 0xF1)  # Indigo
            h.paragraph_format.space_before = Pt(8)

        # Bullet
        elif line.startswith(('- ', '• ', '* ')):
            content = line[2:]
            p = doc.add_paragraph(style='List Bullet')
            _add_inline_formatting(p, content)
            p.paragraph_format.space_after = Pt(3)

        # Numbered list
        elif re.match(r'^\d+\.\s', line):
            content = re.sub(r'^\d+\.\s', '', line)
            p = doc.add_paragraph(style='List Number')
            _add_inline_formatting(p, content)

        # Horizontal rule
        elif line in ('---', '===', '***'):
            _add_horizontal_line(doc)

        # Normal paragraph
        else:
            p = doc.add_paragraph()
            _add_inline_formatting(p, line)
            p.paragraph_format.space_after = Pt(6)

    doc.save(filename)
    print(f"[Word] Saved: {filename}")
    return filename


def _add_inline_formatting(paragraph, text: str):
    """
    Parse **bold**, *italic*, and `code` inline markdown
    and add appropriately formatted runs to the paragraph.
    """
    # Pattern: **bold**, *italic*, `code`
    pattern = r'(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)'
    parts = re.split(pattern, text)
    for part in parts:
        if part is None:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.color.rgb = RGBColor(0x6D, 0x28, 0xD9)
        elif not re.match(r'^\*\*.*\*\*$|^\*.*\*$|^`.*`$', part):
            paragraph.add_run(part)
